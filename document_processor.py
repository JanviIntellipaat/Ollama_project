# document_processor.py
import re
import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional

from docx import Document
try:
    from pypdf import PdfReader
    _HAS_PDF = True
except Exception:
    _HAS_PDF = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Reference extractors ---
EXCEL_REF_RE = re.compile(
    r"""(?ix)
        (?:sheet|tab|tabelle)\s*[:=]?\s*([A-Za-z0-9_ \-]+)?
        .*?
        (?:file|\bfrom\b|source)\s*[:=]?\s*([A-Za-z0-9_\-\.]+\.(?:xlsx|xls|csv))
    """
)
TABLE_HINT_RE = re.compile(r"(?i)\btable\b|\btab\.|\bsheet\b")
FIELD_TAG_RE = re.compile(r"\[(\d+)\]\s*([-A-Za-z0-9_ ]+)")

def _extract_external_refs_from_text(text: str):
    refs = []
    for m in EXCEL_REF_RE.finditer(text):
        sheet_or_table = (m.group(1) or "").strip()
        filename = m.group(2).strip()
        refs.append({
            "type": "excel",
            "filename_hint": filename,
            "sheet_hint": sheet_or_table or None
        })
    for m in TABLE_HINT_RE.finditer(text):
        refs.append({"type": "hint", "table_hint": m.group(0)})
    for m in FIELD_TAG_RE.finditer(text):
        refs.append({"type": "field_ref", "field_no": m.group(1), "field_name": m.group(2).strip()})
    # dedupe
    dedup = []
    seen = set()
    for r in refs:
        k = json.dumps(r, sort_keys=True)
        if k not in seen:
            seen.add(k); dedup.append(r)
    return dedup

def _chunk_text(text: str, max_chars: int = 1200, overlap: int = 100) -> List[str]:
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    if max_chars <= 0:
        max_chars = 1200
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start = end - overlap
        if start < 0:
            start = 0
    return chunks

class DocumentProcessor:
    """
    Process unstructured documents (.docx, .pdf) into chunks with metadata.
    """

    # Back-compat: existing calls might use .process_document; keep it working.
    def process_document(self, file_path: str) -> Dict[str, Any]:
        return self.process(file_path)

    def process(self, file_path: str) -> Dict[str, Any]:
        ext = Path(file_path).suffix.lower()
        if ext == ".docx":
            return self._process_docx(file_path)
        if ext == ".pdf":
            if not _HAS_PDF:
                raise RuntimeError("PDF support requires 'pypdf'. Install it and retry.")
            return self._process_pdf(file_path)
        # Fallback: treat as plain text
        return self._process_text(Path(file_path).read_text(encoding="utf-8", errors="ignore"),
                                  filename=Path(file_path).name,
                                  file_type=ext.lstrip("."))

    def _process_text(self, text: str, filename: str, file_type: str) -> Dict[str, Any]:
        processed_at = datetime.now().isoformat()
        chunks = []
        for ch in _chunk_text(text):
            refs = _extract_external_refs_from_text(ch)
            chunks.append({
                "text": ch,
                "word_count": len(ch.split()),
                "char_count": len(ch),
                "section_heading": "",
                "from_table": False,
                "table_headers": [],
                "external_refs": json.dumps(refs, ensure_ascii=False)
            })
        return {
            "chunks": chunks,
            "tables": [],
            "metadata": {
                "filename": filename,
                "file_type": file_type,
                "processed_at": processed_at
            }
        }

    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            try:
                texts.append(page.extract_text() or "")
            except Exception:
                texts.append("")
        full_text = "\n\n".join(texts)
        return self._process_text(full_text, filename=Path(file_path).name, file_type="pdf")

    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        doc = Document(file_path)
        filename = Path(file_path).name
        processed_at = datetime.now().isoformat()

        chunks: List[Dict[str, Any]] = []
        tables_meta: List[Dict[str, Any]] = []

        current_heading: Optional[str] = None
        para_buffer: List[str] = []

        def _flush_para_buffer():
            nonlocal para_buffer, current_heading, chunks
            if para_buffer:
                full = "\n".join(para_buffer).strip()
                for ch in _chunk_text(full):
                    refs = _extract_external_refs_from_text(ch)
                    chunks.append({
                        "text": ch,
                        "word_count": len(ch.split()),
                        "char_count": len(ch),
                        "section_heading": current_heading or "",
                        "from_table": False,
                        "table_headers": [],
                        "external_refs": json.dumps(refs, ensure_ascii=False)
                    })
                para_buffer = []

        # Paragraphs / headings
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if not text:
                continue
            style = (getattr(p.style, "name", "") or "").lower()
            if "heading" in style:
                _flush_para_buffer()
                current_heading = text
            else:
                para_buffer.append(text)
        _flush_para_buffer()

        # Tables
        for t in doc.tables:
            headers: List[str] = []
            rows: List[List[str]] = []
            for row in t.rows:
                values = [cell.text.strip() for cell in row.cells]
                rows.append(values)
            if rows:
                headers = rows[0]
            # Include a small markdown preview of the table
            md_lines = []
            if headers:
                md_lines.append("| " + " | ".join(h or "" for h in headers) + " |")
                md_lines.append("| " + " | ".join("---" for _ in headers) + " |")
                for row in rows[1:6]:
                    md_lines.append("| " + " | ".join(row) + " |")
            table_md = "\n".join(md_lines) if md_lines else ""
            if table_md:
                refs = _extract_external_refs_from_text(table_md)
                chunks.append({
                    "text": f"Table:\n{table_md}",
                    "word_count": len(table_md.split()),
                    "char_count": len(table_md),
                    "section_heading": current_heading or "Table",
                    "from_table": True,
                    "table_headers": headers,
                    "external_refs": json.dumps(refs, ensure_ascii=False)
                })
                tables_meta.append({"headers": headers, "rows_preview": rows[1:6]})

        return {
            "chunks": chunks,
            "tables": tables_meta,
            "metadata": {
                "filename": filename,
                "file_type": "docx",
                "processed_at": processed_at
            }
        }
